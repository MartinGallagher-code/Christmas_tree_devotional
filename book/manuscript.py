#!/usr/bin/env python3
"""
manuscript.py — edit the book in Word, sync the changes back into book.html.

Two commands:

    ./manuscript.py extract     book.html  ->  manuscript.docx
    ./manuscript.py sync        manuscript.docx  ->  book.html  (changed text only)

Workflow
--------
1.  Run `extract` to (re)generate manuscript.docx from book.html.
2.  Open manuscript.docx in Word / Pages / Google Docs and edit the words.
    - Edit freely: paragraphs, headings, reflection questions, scripture
      descriptions, and the scripture text itself.
    - Italic and bold survive the round-trip (they become <em>/<strong>).
    - Verse numbers appear as small superscript markers — leave them in place
      to keep the numbering; the text around them is yours to change.
    - Don't reorder or delete whole blocks. If you want to add or remove a
      section, do that small bit in book.html (or ask Claude) — everything
      else can be done here.
3.  Save the .docx and run `sync`. Only the blocks whose words actually
    changed are rewritten in book.html; every other byte is left untouched,
    so the print layout is safe.
4.  Rebuild the PDF:  ./build.sh

Each editable block carries an invisible Word bookmark (b0001, b0002, …) that
ties it to the matching element in book.html, so syncing is reliable even
though the file is plain prose to you.

Dependency: python-docx (auto-installed on first run if missing).
"""

import html
import os
import re
import subprocess
import sys
from html.parser import HTMLParser

HERE = os.path.dirname(os.path.abspath(__file__))
BOOK_HTML = os.path.join(HERE, "book.html")
DOCX = os.path.join(HERE, "manuscript.docx")


# --------------------------------------------------------------------------- #
#  dependency bootstrap                                                        #
# --------------------------------------------------------------------------- #
def _ensure_docx():
    try:
        import docx  # noqa: F401
    except ImportError:
        print("Installing python-docx …", file=sys.stderr)
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", "--quiet", "python-docx"]
        )


# --------------------------------------------------------------------------- #
#  parsing book.html into editable "leaf" blocks                              #
# --------------------------------------------------------------------------- #
#
#  A *leaf* is a text-bearing element the author edits. Everything else
#  (sections, boxes, lists) is structure we recurse through. Each leaf records
#  the byte offsets of its inner HTML so `sync` can replace just that span.

HEADINGS = {"h1", "h2", "h3", "h4", "h5"}
# spans that are themselves editable leaves (not inline bits inside a leaf)
LEAF_SPAN_CLASSES = {"scripture-ref", "scripture-desc"}


class Leaf:
    __slots__ = (
        "index", "tag", "classes", "kind", "section", "running",
        "inner_start", "inner_end", "inner_html", "list_ordered",
    )

    def __init__(self, **kw):
        for k in self.__slots__:
            setattr(self, k, kw.get(k))


class BookParser(HTMLParser):
    """Walk book.html, recording every editable leaf with exact inner offsets."""

    def __init__(self, raw):
        super().__init__(convert_charrefs=False)
        self.raw = raw
        # absolute offset of the start of each line
        self._line_starts = [0]
        for line in raw.splitlines(keepends=True):
            self._line_starts.append(self._line_starts[-1] + len(line))
        self.stack = []            # list of (tag, classlist)
        self.leaves = []
        # state while inside a leaf (we slice raw, ignoring inner structure)
        self._leaf = None          # dict or None
        self._leaf_depth = 0       # nesting depth of the leaf's own tag name

    # -- offset helpers ----------------------------------------------------- #
    def _abs(self):
        line, col = self.getpos()
        return self._line_starts[line - 1] + col

    # -- context helpers ---------------------------------------------------- #
    def _classes(self, attrs):
        for k, v in attrs:
            if k == "class":
                return (v or "").split()
        return []

    def _attr(self, attrs, name):
        for k, v in attrs:
            if k == name:
                return v
        return None

    def _cur_section(self):
        for tag, classes, _ in reversed(self.stack):
            if tag == "section":
                return classes
        return []

    def _cur_running(self):
        for tag, _, running in reversed(self.stack):
            if tag == "section" and running:
                return running
        return None

    def _nearest_list(self):
        for tag, classes, _ in reversed(self.stack):
            if tag in ("ul", "ol"):
                return tag, classes
        return None, []

    # -- leaf classification ------------------------------------------------ #
    def _is_leaf(self, tag, classes):
        if tag in HEADINGS or tag in ("p", "blockquote", "cite"):
            return True
        if tag == "li":
            ltag, lclasses = self._nearest_list()
            return not (ltag == "ul" and "scripture-list" in lclasses)
        if tag == "span":
            return any(c in LEAF_SPAN_CLASSES for c in classes)
        if tag == "div":
            return "scripture-text" in classes
        return False

    def _kind(self, tag, classes):
        cset = set(classes)
        if tag == "div" and "scripture-text" in cset:
            return "scripture-text"
        if tag == "span" and "scripture-ref" in cset:
            return "scripture-ref"
        if tag == "span" and "scripture-desc" in cset:
            return "scripture-desc"
        if tag == "p":
            for c in ("lead", "eyebrow", "chapter-eyebrow", "part-eyebrow",
                      "session-meta", "session-aim", "session-close",
                      "subtitle", "author", "imprint", "part-verse"):
                if c in cset:
                    return c
            return "p"
        if tag == "li":
            section = self._cur_section()
            if "contents-page" in section:
                if "toc-part" in cset:
                    return "toc-part"
                if "toc-sub" in cset:
                    return "toc-sub"
                return "toc-item"
            return "li"
        if tag in HEADINGS:
            return tag
        if tag == "blockquote":
            return "blockquote"
        if tag == "cite":
            return "cite"
        return tag

    # -- HTMLParser callbacks ----------------------------------------------- #
    def handle_starttag(self, tag, attrs):
        if self._leaf is not None:
            # inside a leaf: only track depth of the leaf's own tag name
            if tag == self._leaf["tag"]:
                self._leaf_depth += 1
            return

        classes = self._classes(attrs)
        if self._is_leaf(tag, classes):
            start_text = self.get_starttag_text() or ""
            inner_start = self._abs() + len(start_text)
            ltag, lclasses = self._nearest_list()
            self._leaf = {
                "tag": tag,
                "classes": classes,
                "kind": self._kind(tag, classes),
                "section": list(self._cur_section()),
                "running": self._cur_running(),
                "inner_start": inner_start,
                "list_ordered": ltag == "ol",
            }
            self._leaf_depth = 0
        else:
            running = self._attr(attrs, "data-running")
            self.stack.append((tag, classes, running))

    def handle_startendtag(self, tag, attrs):
        # void/self-closed (e.g. <img/>, <br/>) — never a leaf container
        if self._leaf is None and not self._is_leaf(tag, attrs):
            pass  # ignore

    def handle_endtag(self, tag):
        if self._leaf is not None:
            if tag == self._leaf["tag"]:
                if self._leaf_depth > 0:
                    self._leaf_depth -= 1
                    return
                inner_end = self._abs()
                lf = self._leaf
                self.leaves.append(Leaf(
                    index=len(self.leaves) + 1,
                    tag=lf["tag"], classes=lf["classes"], kind=lf["kind"],
                    section=lf["section"], running=lf["running"],
                    inner_start=lf["inner_start"], inner_end=inner_end,
                    inner_html=self.raw[lf["inner_start"]:inner_end],
                    list_ordered=lf["list_ordered"],
                ))
                self._leaf = None
            return
        # pop matching frame
        for i in range(len(self.stack) - 1, -1, -1):
            if self.stack[i][0] == tag:
                del self.stack[i:]
                break


def parse_book(raw):
    p = BookParser(raw)
    p.feed(raw)
    return p.leaves


# --------------------------------------------------------------------------- #
#  inline HTML  <->  rich runs                                                #
# --------------------------------------------------------------------------- #
#
#  A "run" is (text, italic, bold, kind) where kind is None or "verse".
#  We only model the inline vocabulary the book actually uses:
#  <em>, <strong>, <br>, and <span class="verse-num">N</span>.

_INLINE_RE = re.compile(
    r"<(/?)(em|strong)\b[^>]*>"
    r"|<br\s*/?>"
    r"|<span class=\"verse-num\">(.*?)</span>"
    r"|<span class=\"session-time\">(.*?)</span>"
    r"|<span class=\"part-verse-ref\">(.*?)</span>"
    r"|<span class=\"toc-page-num\">(.*?)</span>",
    re.DOTALL,
)

# tail spans we lift out of a leaf so the main text stays clean
TAIL_SPANS = {"session-time", "part-verse-ref", "toc-page-num"}


def inner_to_runs(inner):
    """Parse a leaf's inner HTML into (runs, tail) where tail captures a
    trailing locked span (session-time / part-verse-ref / toc-page-num)."""
    runs = []
    tail = None
    italic = bold = False
    pos = 0

    def emit_text(t):
        if not t:
            return
        runs.append({"text": html.unescape(t), "italic": italic,
                     "bold": bold, "kind": None})

    for m in _INLINE_RE.finditer(inner):
        emit_text(inner[pos:m.start()])
        pos = m.end()
        whole = m.group(0)
        if m.group(2):  # em / strong
            closing = m.group(1) == "/"
            if m.group(2) == "em":
                italic = not closing
            else:
                bold = not closing
        elif whole.startswith("<br"):
            runs.append({"text": "\n", "italic": italic, "bold": bold,
                         "kind": "br"})
        elif m.group(3) is not None:  # verse-num
            runs.append({"text": html.unescape(m.group(3)), "italic": False,
                         "bold": False, "kind": "verse"})
        elif m.group(4) is not None:  # session-time
            tail = ("session-time", html.unescape(m.group(4)))
        elif m.group(5) is not None:  # part-verse-ref
            tail = ("part-verse-ref", html.unescape(m.group(5)))
        elif m.group(6) is not None:  # toc-page-num
            tail = ("toc-page-num", html.unescape(m.group(6)))
    emit_text(inner[pos:])
    return runs, tail


# --- house-style entity encoding ------------------------------------------- #
def encode_text(s):
    """Encode plain text to the book's HTML house style (straight quotes as
    entities, em/en dashes, ellipsis, &, <, >)."""
    s = s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    # normalise Word's smart punctuation back to the book's straight-quote
    # prose style; keep dashes / ellipsis / middot as named entities
    s = (s.replace("“", '"').replace("”", '"')
           .replace("‘", "'").replace("’", "'")
           .replace("—", "&mdash;").replace("–", "&ndash;")
           .replace("…", "&hellip;").replace("·", "&middot;")
           .replace(" ", "&nbsp;"))
    return s


def runs_to_inner(runs, tail, kind):
    """Rebuild a leaf's inner HTML from edited rich runs."""
    out = []
    italic = bold = False

    def close_to(target_i, target_b):
        nonlocal italic, bold
        # close in reverse-open order
        if bold and not target_b:
            out.append("</strong>")
            bold = False
        if italic and not target_i:
            out.append("</em>")
            italic = False

    def open_to(target_i, target_b):
        nonlocal italic, bold
        if target_i and not italic:
            out.append("<em>")
            italic = True
        if target_b and not bold:
            out.append("<strong>")
            bold = True

    for r in runs:
        if r["kind"] == "verse":
            close_to(False, False)
            out.append('<span class="verse-num">%s</span>'
                       % encode_text(r["text"]))
            continue
        if r["kind"] == "br":
            close_to(False, False)
            out.append("<br>")
            continue
        ri, rb = bool(r["italic"]), bool(r["bold"])
        close_to(ri, rb)
        open_to(ri, rb)
        out.append(encode_text(r["text"]))
    close_to(False, False)
    inner = "".join(out)

    # re-attach a trailing locked span
    if tail:
        cls, val = tail
        inner += '<span class="%s">%s</span>' % (cls, encode_text(val))
    return inner


# --------------------------------------------------------------------------- #
#  canonical text (for change detection)                                      #
# --------------------------------------------------------------------------- #
def canonical(runs):
    """Whitespace/punctuation-normalised plain text for comparing two versions
    of a block, ignoring entity vs unicode and smart vs straight quotes."""
    parts = []
    for r in runs:
        if r["kind"] == "br":
            parts.append("\n")
        else:
            parts.append(r["text"])
    s = "".join(parts)
    s = (s.replace("“", '"').replace("”", '"')
           .replace("‘", "'").replace("’", "'")
           .replace("—", "--").replace("–", "-")
           .replace("…", "...").replace(" ", " "))
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _canon_str(s):
    """canonical() for a plain HTML string (e.g. a data-running value)."""
    return canonical([{"text": html.unescape(s), "italic": False,
                       "bold": False, "kind": None}])


# --------------------------------------------------------------------------- #
#  EXTRACT : book.html -> manuscript.docx                                     #
# --------------------------------------------------------------------------- #
def _add_bookmark(paragraph, name):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(abs(hash(name)) % 1000000))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), start.get(qn("w:id")))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _runs_into_paragraph(paragraph, runs):
    from docx.shared import Pt
    for r in runs:
        if r["kind"] == "br":
            paragraph.add_run().add_break()
            continue
        run = paragraph.add_run(r["text"])
        run.italic = bool(r["italic"])
        run.bold = bool(r["bold"])
        if r["kind"] == "verse":
            run.font.superscript = True
            run.font.size = Pt(8)


# Paragraph style per leaf kind. Visual emphasis (italic/bold/colour) lives in
# the STYLE, never on the runs — so `sync` never mistakes a block's cosmetic
# styling for an inline <em>/<strong>.
STYLE_FOR_KIND = {
    "h1": "Heading 1", "h2": "Heading 2", "h3": "Heading 3",
    "h4": "Heading 4", "h5": "Heading 5",
    "blockquote": "Intense Quote", "cite": "Caption",
    "scripture-ref": "BookRef", "scripture-desc": "BookDesc",
    "lead": "BookLead",
    "eyebrow": "BookEyebrow", "chapter-eyebrow": "BookEyebrow",
    "part-eyebrow": "BookEyebrow", "session-meta": "BookEyebrow",
    "session-aim": "BookAim", "session-close": "BookAim",
}

# (name -> font spec) custom paragraph styles created in the document
CUSTOM_STYLES = {
    "BookEyebrow": dict(bold=True, color=(0x80, 0x70, 0x55), size=9, caps=True),
    "BookLead":    dict(size=13),
    "BookAim":     dict(italic=True, color=(0x6A, 0x5A, 0x44)),
    "BookRef":     dict(bold=True, color=(0x55, 0x44, 0x2A)),
    "BookDesc":    dict(italic=True, color=(0x80, 0x70, 0x55)),
}


def _ensure_styles(doc):
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Pt, RGBColor
    existing = {s.name for s in doc.styles}
    for name, spec in CUSTOM_STYLES.items():
        if name in existing:
            st = doc.styles[name]
        else:
            st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            st.base_style = doc.styles["Normal"]
        f = st.font
        if "bold" in spec:
            f.bold = spec["bold"]
        if "italic" in spec:
            f.italic = spec["italic"]
        if "size" in spec:
            f.size = Pt(spec["size"])
        if "color" in spec:
            f.color.rgb = RGBColor(*spec["color"])
        if spec.get("caps"):
            f.small_caps = True


def extract():
    _ensure_docx()
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    raw = open(BOOK_HTML, encoding="utf-8").read()
    leaves = parse_book(raw)

    doc = docx.Document()
    _ensure_styles(doc)
    title = doc.add_paragraph("The Tree Tells a Story")
    title.style = doc.styles["Title"]
    sub = doc.add_paragraph("Editable manuscript — edit the words, then run "
                            "`./manuscript.py sync`")
    sub.style = doc.styles["Subtitle"]
    note = doc.add_paragraph(
        "How to edit: change any words you like. Italic and bold are kept. "
        "Small superscript numbers are verse numbers — leave them in place. "
        "Please don't reorder or delete whole blocks; ask Claude for that. "
        "Save and run sync, then ./build.sh to rebuild the PDF."
    )
    for r in note.runs:
        r.italic = True
        r.font.color.rgb = RGBColor(0x80, 0x70, 0x55)

    last_running = object()
    for lf in leaves:
        runs, tail = inner_to_runs(lf.inner_html)

        # section divider when the running header changes
        if lf.running != last_running and lf.running:
            last_running = lf.running
            div = doc.add_paragraph()
            div.alignment = WD_ALIGN_PARAGRAPH.CENTER
            dr = div.add_run("— %s —" % html.unescape(lf.running))
            dr.bold = True
            dr.font.color.rgb = RGBColor(0xA0, 0x90, 0x70)

        p = doc.add_paragraph()
        style = STYLE_FOR_KIND.get(lf.kind)
        if style and style in doc.styles:
            p.style = doc.styles[style]
        elif lf.kind in ("toc-item", "toc-sub", "li") and lf.list_ordered:
            p.style = doc.styles["List Number"]
        elif lf.kind in ("toc-item", "toc-sub", "toc-part", "li"):
            p.style = doc.styles["List Bullet"]

        _runs_into_paragraph(p, runs)

        # show a locked tail (session time / verse ref / page number) greyed.
        # sync strips this echo before reading the block back.
        if tail:
            cls, val = tail
            sep = p.add_run("    " + val)
            sep.font.color.rgb = RGBColor(0xA0, 0x90, 0x70)

        _add_bookmark(p, "b%04d" % lf.index)

    doc.save(DOCX)
    print("Wrote %s  (%d editable blocks)" % (os.path.relpath(DOCX), len(leaves)))


# --------------------------------------------------------------------------- #
#  SYNC : manuscript.docx -> book.html (changed blocks only)                  #
# --------------------------------------------------------------------------- #
def _para_bookmarks(paragraph):
    from docx.oxml.ns import qn
    return [bs.get(qn("w:name"))
            for bs in paragraph._p.findall(qn("w:bookmarkStart"))]


def _para_to_runs(paragraph):
    """Read a Word paragraph back into our run model."""
    from docx.shared import Pt
    runs = []
    for r in paragraph.runs:
        text = r.text
        if not text and not r._element.findall(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br"):
            continue
        is_verse = bool(r.font.superscript)
        # split out hard line-breaks inside a run
        if "\n" in text:
            segs = text.split("\n")
            for i, seg in enumerate(segs):
                if seg:
                    runs.append({"text": seg, "italic": bool(r.italic),
                                 "bold": bool(r.bold),
                                 "kind": "verse" if is_verse else None})
                if i < len(segs) - 1:
                    runs.append({"text": "\n", "italic": False,
                                 "bold": False, "kind": "br"})
        elif text:
            runs.append({"text": text, "italic": bool(r.italic),
                         "bold": bool(r.bold),
                         "kind": "verse" if is_verse else None})
    return runs


def sync():
    _ensure_docx()
    import docx

    raw = open(BOOK_HTML, encoding="utf-8").read()
    leaves = {lf.index: lf for lf in parse_book(raw)}

    doc = docx.Document(DOCX)
    edits = []          # (leaf, new_inner)
    warnings = []       # advisory notes (e.g. running-header drift)

    for para in doc.paragraphs:
        names = [n for n in _para_bookmarks(para)
                 if n and re.fullmatch(r"b\d{4}", n)]
        if not names:
            continue
        idx = int(names[0][1:])
        lf = leaves.get(idx)
        if lf is None:
            continue

        old_runs, tail = inner_to_runs(lf.inner_html)
        # the locked tail isn't represented as runs in the doc body, but the
        # greyed display text is — strip it for comparison
        new_runs = _para_to_runs(para)
        if tail:
            # drop the trailing greyed display run(s) that echo the tail value
            tail_text = tail[1]
            joined = "".join(r["text"] for r in new_runs)
            cut = joined.rfind(tail_text)
            if cut != -1:
                # rebuild new_runs without the tail echo
                kept, count = [], 0
                for r in new_runs:
                    if count >= cut:
                        break
                    seg = r["text"]
                    if count + len(seg) <= cut:
                        kept.append(r)
                        count += len(seg)
                    else:
                        r = dict(r)
                        r["text"] = seg[: cut - count]
                        kept.append(r)
                        count = cut
                new_runs = kept
            # strip trailing whitespace-only run text
            while new_runs and new_runs[-1]["text"].strip() == "" \
                    and new_runs[-1]["kind"] is None:
                new_runs.pop()

        if canonical(old_runs) == canonical(new_runs):
            continue  # unchanged

        new_inner = runs_to_inner(new_runs, tail, lf.kind)
        edits.append((lf, new_inner))

        # warn when a heading that drives a running header changed: the
        # data-running attribute must be updated by hand to match.
        if lf.kind in HEADINGS and lf.running and \
                canonical(old_runs) == _canon_str(lf.running):
            warnings.append(
                'block %d: heading changed but the page running-header '
                '(data-running="%s") still shows the old title. Update '
                "data-running in book.html if you want them to match."
                % (lf.index, html.unescape(lf.running))
            )

    if not edits:
        print("No changes detected.")
        if warnings:
            print("\nNote:")
            for w in warnings:
                print("  ! " + w)
        return

    # apply from the end of the file backwards so offsets stay valid
    edits.sort(key=lambda e: e[0].inner_start, reverse=True)
    out = raw
    for lf, new_inner in edits:
        out = out[:lf.inner_start] + new_inner + out[lf.inner_end:]

    open(BOOK_HTML, "w", encoding="utf-8").write(out)

    print("Updated %d block(s) in %s:" % (len(edits), os.path.relpath(BOOK_HTML)))
    for lf, _ in sorted(edits, key=lambda e: e[0].index):
        ctx = lf.running or " ".join(lf.section)
        print("  - block %d  (%s, %s)" % (lf.index, lf.kind, ctx))
    if warnings:
        print("\nNote:")
        for w in warnings:
            print("  ! " + w)
    print("\nNext: ./build.sh   to rebuild the PDF.")


# --------------------------------------------------------------------------- #
def main():
    cmd = sys.argv[1] if len(sys.argv) > 1 else ""
    if cmd == "extract":
        extract()
    elif cmd == "sync":
        sync()
    else:
        print(__doc__)
        sys.exit(0 if cmd in ("-h", "--help", "help") else 2)


if __name__ == "__main__":
    main()
